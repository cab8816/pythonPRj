from django.http import HttpResponse, FileResponse
from django.shortcuts import render, redirect
from Myapp.models import Students, Pingshenxxb

from docx import Document  # 导入库
from docxtpl import DocxTemplate, InlineImage  # 导入模板库
from docx.shared import Inches, Cm  # 支持修改文字大小的库
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENTATION

from Myapp.models import Biao4
from django.core.paginator import Paginator, Page, PageNotAnInteger, EmptyPage
from docx.shared import Inches

import os


# git@github.com:cab8816/pythonPRJ.git

def testdbadd(request):
    student = Students(name='谢振乾')
    student.save()
    return HttpResponse("<p> 数据添加成功！</p>")


def testdblst(request):
    response1 = ""
    list = Students.objects.all()
    for var in list:
        response1 += var.name + " "
    return HttpResponse("<p>" + response1 + "</p>")


def testdbupd(request):
    stu = Students.objects.get(id=2)
    stu.name = '谢岸马'
    stu.save()
    stu.objects.filter(id=3).update(name='谢多多')
    return HttpResponse("<p>成功修改</p>")


def testdoc(request):
    document = Document()
    document.add_heading('Document Title', 0)
    p = document.add_paragraph('A plain pragraph having some')
    p.add_run('bold').bold = True
    p.add_run('and some')
    p.add_run('italic.').italic = True
    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')
    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )
    # document.add_picture('monty-truth.png', width=Inches(1.25))
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
    document.add_page_break()
    document.save('testdoc/demo.docx')
    return HttpResponse("<p>doc文件生成成功</p>")


def readword(request):
    document = Document('testdoc/psbg.docx')
    table = document.tables[0]
    psxxb = Pingshenxxb.objects.get(id=5)
    for row in table.rows:
        rowcells = row.cells
        biao = Biao4(
            psxxb=psxxb,
            lyxh=rowcells[0].text,
            lyname=rowcells[1].text,
            lbxh=rowcells[2].text,
            lb=rowcells[3].text,
            dxxh=rowcells[4].text,
            duixiang=rowcells[5].text,
            xmxh=rowcells[6].text,
            csmc=rowcells[7].text,
            yjbz=rowcells[8].text,
            xzfw=rowcells[9].text,
            sm=rowcells[10].text,
        )
        biao.save()
    context = {}
    context['msg'] = '读取word文档成功'
    return render(request, 'test.html', context)


def split_page(object_list, request, per_page=8):
    paginator = Paginator(object_list, per_page)
    # 取出当前需要展示的页码, 默认为1
    page_num = request.GET.get('page', default='1')
    # 根据页码从分页器中取出对应页的数据
    try:
        page = paginator.page(page_num)
    except PageNotAnInteger as e:
        # 不是整数返回第一页数据
        # page = paginator.page('1')
        page = paginator.page('1')
        page_num = 1
    except EmptyPage as e:
        # 当参数页码大于或小于页码范围时,会触发该异常
        print('EmptyPage:{}'.format(e))
        if int(page_num) > paginator.num_pages:
            # 大于 获取最后一页数据返回
            page = paginator.page(paginator.num_pages)
        else:
            # 小于 获取第一页
            page = paginator.page(1)

    # 这部分是为了再有大量数据时，仍然保证所显示的页码数量不超过10，
    page_num = int(page_num)
    if page_num < 6:
        if paginator.num_pages <= 10:
            dis_range = range(1, paginator.num_pages + 1)
        else:
            dis_range = range(1, 11)
    elif (page_num >= 6) and (page_num <= paginator.num_pages - 5):
        dis_range = range(page_num - 5, page_num + 5)
    else:
        dis_range = range(paginator.num_pages - 9, paginator.num_pages + 1)

    data = {'page': page, 'paginator': paginator, 'dis_range': dis_range}
    return data


def lstbiao4(request):
    # mlstbiao4 = Biao4.objects.all()[:5]
    mlstbiao4 = Biao4.objects.all()

    data1 = split_page(mlstbiao4, request, 10)
    return render(request, 'biao4.html', data1)


# def cellformat(row_cells,col=0,width=2):
#     row_cells[col].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     row_cells[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
#     row_cells[col].width = Inches(width)

def genbiao4(request):
    # install python-docx https://python-docx.readthedocs.io/en/latest/index.html
    document = Document('testdoc/b4t.docx')  # 新建空文档
    # 设置正文颜色，大小，粗体
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    document.styles['Normal'].font.size = Pt(10)
    document.styles['Normal'].font.bold = False

    # 设置页边距
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    section = document.sections[0]
    # 增加页码
    document.settings.odd_and_even_pages_header_footer = False
    section.even_page_header.is_linked_to_previous

    section.orientation = WD_ORIENTATION.PORTRAIT  # LANDSCAPE
    # header = section.header
    # paragraph = header.paragraphs[0]
    # paragraph.text = "title of my document"
    #  header.is_linked_to_previous = True

    # footer = section.footer
    # p1 = footer.paragraphs[0]
    # p1.text = footer.pageNumber.text

    # document.add_heading('4表   建议批准的检验检测能力表',2) #增加标题“Document Title”，第二个参数“0”表示是标题
    # p = document.add_paragraph()

    # p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # run = p.add_run('4表')
    # font = run.font
    # font.name = u'黑体'
    # # document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    # font.size = Pt(13)
    # p = document.add_paragraph()
    # p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # run = p.add_run('建议批准的检验检测能力表')
    # font = run.font
    # font.name = '黑体'
    #
    # font.size = Pt(13)
    # font.bold = True
    #
    # p = document.add_paragraph()
    # p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # run = p.add_run('检验检测场所地址:')
    # font = run.font
    # font.name = '黑体'
    # font.size = Pt(13)
    # font.italic = True
    #
    # p.add_run('广州市南沙区东涌镇市南公路东涌段115号')
    # mlstbiao4 = Biao4.objects.all()[:50]
    mlstbiao4 = Biao4.objects.all()
    table = document.add_table(rows=1, cols=11, style='Table Grid')
    table.width = Inches(200)
    table.autofit = True

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '领域序号'
    hdr_cells[1].text = '领域'
    hdr_cells[2].text = '类别序号'
    hdr_cells[3].text = '类别'
    hdr_cells[4].text = '对象序号'
    hdr_cells[5].text = '检测对象'
    hdr_cells[6].text = '项目/参数'
    # hdr_cells[7].text = ''
    hdr_cells[8].text = '依据的标准（方法）名称及编号（含年号）'
    hdr_cells[9].text = '限制范围'
    hdr_cells[10].text = '说明'
    hdr_cells = table.add_row().cells
    # hdr_cells[0].text = ''
    # hdr_cells[1].text = ''
    # hdr_cells[2].text = ''
    # hdr_cells[3].text = ''
    # hdr_cells[4].text = ''
    # hdr_cells[5].text = ''
    hdr_cells[6].text = '序号'
    hdr_cells[7].text = '名称'
    # hdr_cells[8].text = ''
    # hdr_cells[9].text = ''
    # hdr_cells[10].text = ''
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(1, 1))
    table.cell(0, 2).merge(table.cell(1, 2))
    table.cell(0, 3).merge(table.cell(1, 3))
    table.cell(0, 4).merge(table.cell(1, 4))
    table.cell(0, 5).merge(table.cell(1, 5))
    table.cell(0, 6).merge(table.cell(0, 7))
    table.cell(0, 8).merge(table.cell(1, 8))
    table.cell(0, 9).merge(table.cell(1, 9))
    table.cell(0, 10).merge(table.cell(1, 10))
    for cell in table.rows[0].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for cell in table.rows[1].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in mlstbiao4:
        nrow = table.add_row()
        row_cells = nrow.cells
        nrow.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        row_cells[0].text = row.lyxh
        row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        row_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_cells[0].width = Inches(4)

        row_cells[1].text = row.lyname
        row_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        row_cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_cells[1].width = Inches(12)

        row_cells[2].text = row.lbxh
        row_cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        row_cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_cells[2].width = Inches(6)

        row_cells[3].text = row.lb
        row_cells[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        row_cells[3].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_cells[3].width = Inches(12)

        row_cells[4].text = row.dxxh
        row_cells[4].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        row_cells[4].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_cells[4].width = Inches(6)

        row_cells[5].text = row.duixiang
        row_cells[5].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        row_cells[5].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_cells[5].width = Inches(12)

        row_cells[6].text = row.xmxh
        row_cells[6].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        row_cells[6].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_cells[6].width = Inches(5)

        row_cells[7].text = row.csmc
        row_cells[7].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        row_cells[7].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_cells[7].width = Inches(8)

        row_cells[8].text = row.yjbz
        row_cells[8].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        row_cells[8].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_cells[8].width = Inches(12)

        row_cells[9].text = row.xzfw
        row_cells[9].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        row_cells[9].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_cells[9].width = Inches(6)

        row_cells[10].text = row.sm
        row_cells[10].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        row_cells[10].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_cells[10].width = Inches(6)

    # document.add_page_break()
    document.save('testdoc/genword.docx')
    context = {}
    context['msg'] = '生成word文档成功'
    return render(request, 'test.html', context)


def changeword(request):
    # 在cmd中使用pip3  install docxtpl
    doc = DocxTemplate('testdoc/changeword.docx')
    content = {'mycontent': 'hello xie an ma'}
    doc.render(content)
    doc.save('testdoc/changeword.docx')

    context = {}
    context['msg'] = '替换word文档内容成功！'
    return render(request, 'test.html', context)


def downloadfile(request):
    context = {}
    file = open('media/baojia.jpg', 'rb')
    response = FileResponse(file)
    response['Content-Type'] = 'application/octet-stream'
    response['Content-Disposition'] = 'attachment;filename="mbaojia.jpg"'
    return response

    return render(request, "downloadfile.html")


def uploadfile(request):
    context = {}
    if request.method == 'POST':
        myFile = request.FILES.get("myfile", None)
        if not myFile:
            return HttpResponse('没有文件来上传！')
        print(myFile.name)
        destination = open(os.path.join("media/", myFile.name), 'wb+')
        for chunk in myFile.chunks():
            destination.write(chunk)
        destination.close()
        context['msg'] = '文件上传完成！'
        # return redirect('text.html',context)

    return render(request, "uploadfile.html", context)
